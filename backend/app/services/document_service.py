from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import os
from typing import List, Dict, Tuple

class DocumentService:
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)

    def parse_document(self, file_path: str) -> Tuple[List[str], Document]:
        """
        Parse a Word document and return its paragraphs and the document object
        """
        try:
            print(f"Opening document at path: {file_path}")
            doc = Document(file_path)
            print(f"Successfully opened document. Number of paragraphs: {len(doc.paragraphs)}")
            paragraphs = [p.text for p in doc.paragraphs]
            print(f"Extracted {len(paragraphs)} paragraphs")
            return paragraphs, doc
        except Exception as e:
            print(f"Error parsing document: {str(e)}")
            print(f"Error type: {type(e)}")
            raise

    def create_redline_document(self, original_doc: Document, changes: List[Dict]) -> Document:
        """
        Create a redline version of the document with suggested changes
        """
        redline_doc = Document()
        
        # Copy document properties safely
        try:
            if hasattr(original_doc, 'core_properties'):
                redline_doc.core_properties = original_doc.core_properties
        except Exception as e:
            print(f"Warning: Could not copy document properties: {str(e)}")
        
        # Process each paragraph
        for paragraph in original_doc.paragraphs:
            new_paragraph = redline_doc.add_paragraph()
            
            # Copy paragraph formatting
            try:
                new_paragraph.style = paragraph.style
            except Exception as e:
                print(f"Warning: Could not copy paragraph style: {str(e)}")
            
            # Process runs (text with specific formatting)
            for run in paragraph.runs:
                # Check if this run contains any changes
                is_changed = any(
                    change['original_text'] in run.text 
                    for change in changes
                )
                
                if is_changed:
                    # Add strikethrough for original text
                    new_run = new_paragraph.add_run(run.text)
                    new_run.font.strike = True
                    new_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                    
                    # Add suggested text
                    for change in changes:
                        if change['original_text'] in run.text:
                            suggestion = new_paragraph.add_run(f" → {change['suggested_text']}")
                            suggestion.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                else:
                    # Copy original text without changes
                    new_run = new_paragraph.add_run(run.text)
                    try:
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                    except Exception as e:
                        print(f"Warning: Could not copy run formatting: {str(e)}")

        return redline_doc

    def create_clean_document(self, original_doc: Document, changes: List[Dict]) -> Document:
        """
        Create a clean version of the document with changes applied
        """
        clean_doc = Document()
        
        # Copy document properties safely
        try:
            if hasattr(original_doc, 'core_properties'):
                clean_doc.core_properties = original_doc.core_properties
        except Exception as e:
            print(f"Warning: Could not copy document properties: {str(e)}")
        
        # Process each paragraph
        for paragraph in original_doc.paragraphs:
            new_paragraph = clean_doc.add_paragraph()
            try:
                new_paragraph.style = paragraph.style
            except Exception as e:
                print(f"Warning: Could not copy paragraph style: {str(e)}")
            
            # Process runs
            for run in paragraph.runs:
                # Check if this run contains any changes
                is_changed = any(
                    change['original_text'] in run.text 
                    for change in changes
                )
                
                if is_changed:
                    # Apply suggested changes
                    for change in changes:
                        if change['original_text'] in run.text:
                            new_text = run.text.replace(
                                change['original_text'], 
                                change['suggested_text']
                            )
                            new_run = new_paragraph.add_run(new_text)
                            try:
                                new_run.font.name = run.font.name
                                new_run.font.size = run.font.size
                                new_run.font.bold = run.font.bold
                                new_run.font.italic = run.font.italic
                            except Exception as e:
                                print(f"Warning: Could not copy run formatting: {str(e)}")
                else:
                    # Copy original text without changes
                    new_run = new_paragraph.add_run(run.text)
                    try:
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold
                        new_run.font.italic = run.font.italic
                    except Exception as e:
                        print(f"Warning: Could not copy run formatting: {str(e)}")

        return clean_doc

    def save_document(self, doc: Document, file_path: str) -> str:
        """
        Save a document to the specified path
        """
        doc.save(file_path)
        return file_path 