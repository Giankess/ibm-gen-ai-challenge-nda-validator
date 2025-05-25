from docx import Document
from docx.shared import RGBColor
from typing import List, Dict, Tuple
import re
from collections import defaultdict
import os
import subprocess
import tempfile
import zipfile
from lxml import etree
import io

class TrainingAnalyzer:
    def __init__(self, training_dir: str = None):
        if training_dir is None:
            # Get the absolute path to the training_data directory
            current_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            self.training_dir = os.path.join(current_dir, "training_data")
            print(f"Training directory path: {self.training_dir}")  # Debug log
            if not os.path.exists(self.training_dir):
                print(f"Warning: Training directory not found at {self.training_dir}")
                # Try alternative path
                alt_path = os.path.join(os.path.dirname(current_dir), "backend", "training_data")
                print(f"Trying alternative path: {alt_path}")
                if os.path.exists(alt_path):
                    self.training_dir = alt_path
                    print(f"Found training data at: {self.training_dir}")
        else:
            self.training_dir = training_dir
        self.patterns = defaultdict(list)
        self.suggestions = defaultdict(list)
        self.context_patterns = defaultdict(list)

    def analyze_training_data(self) -> Dict:
        """
        Analyze all training documents and extract patterns
        """
        if not os.path.exists(self.training_dir):
            raise FileNotFoundError(f"Training directory {self.training_dir} not found")

        print(f"\nAnalyzing training data in directory: {self.training_dir}")
        print(f"Found files: {os.listdir(self.training_dir)}")

        # Reset patterns before analysis
        self.patterns = defaultdict(list)
        self.suggestions = defaultdict(list)
        self.context_patterns = defaultdict(list)

        for filename in os.listdir(self.training_dir):
            if filename.endswith(('.docx', '.doc')):
                file_path = os.path.join(self.training_dir, filename)
                try:
                    print(f"\nProcessing file: {filename}")
                    self._analyze_document(file_path)
                except Exception as e:
                    print(f"Warning: Could not analyze {filename}: {str(e)}")

        compiled_patterns = self._compile_patterns()
        
        # Print detailed summary of compiled patterns
        print("\nTraining Data Analysis Summary:")
        print("===============================")
        for category, data in compiled_patterns.items():
            print(f"\nCategory: {category}")
            print(f"Number of patterns: {len(data['patterns'])}")
            print(f"Number of suggestions: {len(data['suggestions'])}")
            print(f"Number of context patterns: {len(data['context'])}")
            if data['patterns']:
                print("\nSample patterns:")
                for i, pattern in enumerate(data['patterns'][:3]):  # Show first 3 patterns
                    print(f"\nPattern {i+1}:")
                    print(f"Original: {pattern}")
                    print(f"Suggested: {data['suggestions'][i] if i < len(data['suggestions']) else 'No suggestion'}")
                    if i < len(data['context']):
                        print(f"Context: {data['context'][i]}")

        return compiled_patterns

    def _extract_redline_changes(self, docx_path: str) -> List[Dict]:
        """
        Extract all redline changes from a .docx file using direct XML parsing
        """
        changes = []
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_file:
                # Read document.xml content
                document_xml_content = docx_file.read('word/document.xml')
                
                # Parse the XML content
                root = etree.fromstring(document_xml_content)
                
                # Find all <w:ins> and <w:del> elements
                revision_elements = root.xpath('//w:ins | //w:del', namespaces=namespaces)
                
                print(f"\nFound {len(revision_elements)} revision elements in {docx_path}")
                
                for element in revision_elements:
                    change_type = None
                    if element.tag == f"{{{namespaces['w']}}}ins":
                        change_type = "insertion"
                    elif element.tag == f"{{{namespaces['w']}}}del":
                        change_type = "deletion"
                    else:
                        continue
                    
                    # Get the text content
                    text = "".join(element.xpath('.//w:t/text()', namespaces=namespaces))
                    
                    # Get the author and date if available
                    author = element.get(f"{{{namespaces['w']}}}author", "Unknown")
                    date = element.get(f"{{{namespaces['w']}}}date", "Unknown")
                    
                    print(f"\nFound {change_type}:")
                    print(f"Text: {text}")
                    print(f"Author: {author}")
                    print(f"Date: {date}")
                    
                    changes.append({
                        "type": change_type,
                        "text": text,
                        "author": author,
                        "date": date
                    })
                
        except Exception as e:
            print(f"Error processing {docx_path}: {str(e)}")
        
        return changes

    def _analyze_document(self, file_path: str):
        """
        Analyze a single training document for redline changes
        """
        print(f"\nAnalyzing document: {file_path}")
        
        if file_path.endswith('.doc'):
            # Convert .doc to .docx using LibreOffice
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                docx_path = tmp.name
            try:
                subprocess.run([
                    'soffice',
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', os.path.dirname(docx_path),
                    file_path
                ], check=True)
                doc = Document(docx_path)
            finally:
                if os.path.exists(docx_path):
                    os.unlink(docx_path)
        else:
            docx_path = file_path
            doc = Document(file_path)
        
        # Extract redline changes using XML parsing
        changes = self._extract_redline_changes(docx_path)
        
        # Process the changes to find patterns
        current_deletions = []
        current_insertions = []
        current_context = []
        
        for change in changes:
            if change["type"] == "deletion":
                current_deletions.append(change["text"])
            elif change["type"] == "insertion":
                current_insertions.append(change["text"])
            
            # If we have both deletions and insertions, store the pattern
            if current_deletions and current_insertions:
                print(f"\nFound pattern:")
                print(f"Deletions: {' '.join(current_deletions)}")
                print(f"Insertions: {' '.join(current_insertions)}")
                print(f"Context: {' '.join(current_context)}")
                
                self._store_pattern(current_deletions, current_insertions, current_context)
                current_deletions = []
                current_insertions = []
                current_context = []
        
        # Also analyze the document structure for context
        for paragraph in doc.paragraphs:
            if not any(change["text"] in paragraph.text for change in changes):
                current_context.append(paragraph.text)

    def _store_pattern(self, strikethrough: List[str], red: List[str], context: List[str]):
        """
        Store a pattern from strikethrough and red text
        """
        original_text = " ".join(strikethrough).strip()
        suggested_text = " ".join(red).strip()
        context_text = " ".join(context).strip()
        
        if original_text or suggested_text:  # Changed from 'and' to 'or' since we might have only insertions or deletions
            # Extract the pattern category
            category = self._categorize_pattern(original_text or suggested_text, suggested_text or original_text)
            
            print(f"\nStoring new pattern:")
            print(f"Category: {category}")
            print(f"Original: {original_text}")
            print(f"Suggested: {suggested_text}")
            print(f"Context: {context_text}")
            
            # Store the pattern
            self.patterns[category].append({
                "original": original_text,
                "suggested": suggested_text,
                "context": context_text
            })
            
            # Store the suggestion pattern
            if suggested_text:
                self.suggestions[category].append(suggested_text)
            
            # Store context patterns
            if context_text:
                self.context_patterns[category].append(context_text)
            
            print(f"Current pattern counts for {category}:")
            print(f"Patterns: {len(self.patterns[category])}")
            print(f"Suggestions: {len(self.suggestions[category])}")
            print(f"Context patterns: {len(self.context_patterns[category])}")

    def _categorize_pattern(self, original: str, suggested: str) -> str:
        """
        Categorize the pattern based on content
        """
        # Define category keywords
        categories = {
            "confidentiality": ["confidential", "secret", "proprietary", "trade secret", "disclose", "disclosure"],
            "duration": ["perpetual", "term", "period", "duration", "expiration", "during"],
            "scope": ["scope", "purpose", "use", "application", "all", "any", "business"],
            "liability": ["liability", "damages", "indemnification", "warranty", "warrant"],
            "intellectual_property": ["intellectual property", "ip", "patent", "copyright", "trademark", "license"],
            "assignment": ["assign", "transfer", "convey", "license", "grant"],
            "termination": ["terminate", "termination", "end", "expire"],
            "governing_law": ["governing law", "jurisdiction", "venue", "dispute", "applicable law"]
        }
        
        # Check which category keywords appear in either the original or suggested text
        text_lower = (original + " " + suggested).lower()
        for category, keywords in categories.items():
            if any(keyword in text_lower for keyword in keywords):
                return category
        
        return "other"

    def _extract_common_patterns(self, texts: List[str]) -> List[str]:
        """
        Extract common patterns from a list of texts
        """
        if not texts:
            return []
            
        # Convert texts to lowercase for comparison
        texts_lower = [text.lower() for text in texts]
        
        # Find common single words
        common_words = set()
        word_counts = defaultdict(int)
        for text in texts_lower:
            words = text.split()
            for word in words:
                if len(word) > 3:  # Only consider words longer than 3 characters
                    word_counts[word] += 1
        
        # Add words that appear in at least 2 texts
        for word, count in word_counts.items():
            if count >= 2:
                common_words.add(word)
        
        # Find common phrases (2 or more words)
        common_phrases = set()
        for text in texts_lower:
            words = text.split()
            # Look for phrases of 2-4 words
            for i in range(len(words)):
                for j in range(2, min(5, len(words) - i + 1)):
                    phrase = " ".join(words[i:i+j])
                    if len(phrase) > 5 and all(phrase in t for t in texts_lower):
                        common_phrases.add(phrase)
        
        # Combine single words and phrases
        patterns = list(common_words) + list(common_phrases)
        
        # Sort by length (longer patterns first) and return
        return sorted(patterns, key=len, reverse=True)

    def _compile_patterns(self) -> Dict:
        """
        Compile the analyzed patterns into a structured format
        """
        compiled_patterns = {}
        
        print("\nCompiling patterns:")
        print("==================")
        
        for category, patterns in self.patterns.items():
            print(f"\nCategory: {category}")
            print(f"Number of patterns: {len(patterns)}")
            
            if not patterns:
                continue
            
            # Store all original patterns
            original_patterns = [p["original"] for p in patterns if p["original"]]
            
            # Store all suggestions
            suggestion_patterns = [p["suggested"] for p in patterns if p["suggested"]]
            
            # Store all context patterns
            context_patterns = [p["context"] for p in patterns if p["context"]]
            
            # Only add category if we have patterns
            if original_patterns or suggestion_patterns:
                compiled_patterns[category] = {
                    "patterns": original_patterns,
                    "suggestions": suggestion_patterns,
                    "context": context_patterns,
                    "examples": patterns[:5]  # Keep first 5 examples
                }
                
                print(f"Compiled {len(original_patterns)} patterns")
                print(f"Compiled {len(suggestion_patterns)} suggestions")
                print(f"Compiled {len(context_patterns)} context patterns")
        
        if not compiled_patterns:
            print("No patterns found in training data, using default patterns")
            compiled_patterns = {
                "confidentiality": {
                    "patterns": ["all information", "any information"],
                    "suggestions": ["specifically identified confidential information"],
                    "context": ["confidential", "secret", "proprietary"],
                    "examples": []
                },
                "duration": {
                    "patterns": ["perpetual", "indefinite"],
                    "suggestions": ["for a period of 5 years"],
                    "context": ["confidentiality", "obligation"],
                    "examples": []
                }
            }
        
        return compiled_patterns 